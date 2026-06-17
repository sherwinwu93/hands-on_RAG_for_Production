import os
from docx import Document
from IPython.display import display, Image
import zipfile

def extract_para(document: Document):
  """提取段落文字(没有表格文字)"""
  for para in document.paragraphs:
    if len(para.text.strip()) > 0:  # Only print non-empty paragraphs
      print(para.text, end="\n---\n")
def extract_table(document: Document):
  """提取表格"""
  for table in document.tables:
    print("\n--- Table ---")
    for row in table.rows:
      row_text = [cell.text for cell in row.cells]
      print(row_text)

def extract_image(document: Document):
  """提取图片"""
  zipf = zipfile.ZipFile('sample_data/sample_data.docx')
  filelist = zipf.namelist()
  # print (filelist)

  for fname in filelist:
    _, ext = os.path.splitext(fname)
    if ext in ['.jpg', '.jpeg', '.png', '.gif']:
      # read image and display in Jupyter
      with zipf.open(fname) as img_file:
          img_data = img_file.read()
          display(Image(data=img_data))

document = Document("sample_data/sample_data.docx")
# extract_para(document)
# extract_table(document)
# extract_image(document)